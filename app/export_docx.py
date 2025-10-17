from docx import Document

def export(items, ar_facts, en_facts, path="quiz.docx", show_ar=True):
    doc=Document(); doc.add_heading("English Quiz (Teacher-fed)",0)
    for i,it in enumerate(items,1):
        doc.add_paragraph(f"{i}. {it['stem']}")
        if it['type']=="MCQ":
            ABCD="ABCD"
            for j,c in enumerate(it['choices']):
                doc.add_paragraph(f"   {ABCD[j]}. {c}")
        if show_ar:
            doc.add_paragraph("   Evidence:")
            for fid in sorted(set(it['fact_ids'])):
                doc.add_paragraph(f"     • AR: {ar_facts[fid]}")
                doc.add_paragraph(f"       EN: {en_facts[fid]}")
    doc.add_page_break(); doc.add_heading("Answer Key",1)
    for i,it in enumerate(items,1):
        doc.add_paragraph(f"{i}. {it['answer']}")
    doc.save(path); return path
